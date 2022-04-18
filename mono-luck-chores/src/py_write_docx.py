from pydoc import doc
import docx
from docx.shared import  Pt
from docx.oxml.ns import qn
from docx_convert_pdf import *

def write_docx(lockid,member):
    file = docx.Document("../doc/置物櫃租用申請表.docx")
    tables = file.tables
    table = tables[0]
    membersdata = ["name", "phone", "mail"]
    membersdata[0] = table.cell(1, 1).paragraphs[0].add_run(member["name"])
    membersdata[1] = table.cell(1, 3).paragraphs[0].add_run(member["phone"])
    membersdata[2] = table.cell(2, 1).paragraphs[0].add_run(member["mail"])

    for memberdata in membersdata:
        memberdata.font.name = '微軟正黑體'
        memberdata._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
        memberdata.font.size = Pt(12)
    file.save("../doc/"+lockid+".docx")

def member_dict(data):
    for key,value in data.items():
        if(value!=None):
            write_docx(key,value)
            docx_convert_pdf(key)

